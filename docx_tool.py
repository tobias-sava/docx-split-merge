# DOCX Split / Merge Tool - Created by Tobias Sava

# For a flawless experience, read README.md before using the tool.

import argparse
import os
from docx import Document
from docx.oxml import OxmlElement

def split_docx(file_path, output_prefix='split_part_'):
    
    """Splits a DOCX file at page breaks while preserving the media and format."""
    
    if not os.path.exists(file_path):
        print(f"Error: File '{file_path}' not found.")
        return

    doc = Document(file_path)

    if not doc.paragraphs and not doc.tables: # Handling empty file (if no paragraphs or tables)
        print("Error: The document is empty.")
        return
    
    new_doc = Document() # Creating a new document for storing the split parts.
    split_count = 1 # Keeps tracks of the amount of splits.

    for element in doc.element.body:
        if element.tag.endswith('br'): # Detects page breaks.
            new_doc.save(f"{output_prefix}{split_count}.docx") # Saves current split.
            new_doc = Document() # Starts a new document for the new part.
            split_count += 1
        else:
            new_doc.element.body.append(element) # Copies elements.

    new_doc.save(f"{output_prefix}{split_count}.docx")
    print(f"Splitting complete - {split_count} files were created.")

def merge_docx(file_list, output_file="merged.docx"):
    
    """Merges multiple DOCX files while preversing the media and format."""

    merged_doc = Document()

    for file in file_list:
        if not file.endswith('.docx'): # Only .docx files will be processed.
            print(f"Skipping {file} due to unsupported format.")
            continue

        if not os.path.exists(file): # Skipping non-existant files.
            print(f"Error: File {file} not found. Skipping . . .")
            continue

        doc = Document(file)

        if not doc.paragraphs and not doc.tables:
            print(f"Skipping {file} due to it being empty.")
            continue

        for element in doc.element.body: # Copying elements.
            merged_doc.element.body.append(element)

        merged_doc.add_page_break() # Adding page break between the merged files.

    merged_doc.save(output_file)
    print(f"Documents have been merged into {output_file}.")

def main():
    """Simple CLI for selecting between split / merge modes."""

    parser = argparse.ArgumentParser(description="Split or merge DOCX files.")

    # Defining sub-commands (split / merge).

    subparsers = parser.add_subparsers(dest="command")

    # Split command:

    split_parser = subparsers.add_parser("split", help="Split a DOCX file.")
    split_parser.add_argument("file", help="Path of the DOCX file to split.")

    # Merge command:

    merge_parser = subparsers.add_parser("merge", help="Merge multiple DOCX files into one.")
    merge_parser.add_argument("files", narga="+", help="List of DOCX files to merge.")
    merge_parser.add_argument("--output", default="merged.docx", help="Output file name (default: merged.docx)")

    args = parser.parse_args()

    # Calling functions based on user input:

    if args.command == "split":
        split_docx(args.file)
    elif args.command == "merge":
        merge_docx(args.files, args.output)
    else:
        parser.print_help() # Showing the respective help message if no valid command is given.

if __name__ == "__main__":
    main()